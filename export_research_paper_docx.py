from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pypdf import PdfReader


PROJECT_ROOT = Path(__file__).resolve().parent
PDF_PATH = PROJECT_ROOT / "major_project_research_paper.pdf"
DOCX_PATH = PROJECT_ROOT / "major_project_research_paper.docx"


def clean_text(text: str) -> list[str]:
    lines = []
    for raw_line in text.splitlines():
        line = " ".join(raw_line.split()).strip()
        if not line:
            lines.append("")
            continue
        if line.startswith("Page "):
            continue
        lines.append(line)
    return lines


def add_paragraph(document: Document, text: str) -> None:
    lowered = text.lower()
    if text == "References":
        document.add_page_break()
        p = document.add_paragraph()
        p.style = document.styles["Heading 1"]
        p.add_run(text)
        return
    if text.startswith("Design and Evaluation of"):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        return
    if text.startswith("Research Paper Based on") or text.startswith("Prepared from the"):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        return
    if len(text) < 80 and text[:1].isdigit() and "." in text[:5]:
        p = document.add_paragraph()
        p.style = document.styles["Heading 1"]
        p.add_run(text)
        return
    if len(text) < 90 and text.startswith("Appendix "):
        p = document.add_paragraph()
        p.style = document.styles["Heading 1"]
        p.add_run(text)
        return
    if (
        len(text) < 110
        and not text.endswith(".")
        and not text.endswith(":")
        and not lowered.startswith("keywords")
        and not lowered.startswith("table ")
        and not lowered.startswith("figure ")
        and not text.startswith("[")
    ):
        p = document.add_paragraph()
        p.style = document.styles["Heading 2"]
        p.add_run(text)
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11)


def build_docx() -> None:
    reader = PdfReader(str(PDF_PATH))
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Heading 2"].font.size = Pt(12.5)

    for page_index, page in enumerate(reader.pages):
        lines = clean_text(page.extract_text() or "")
        previous_blank = True
        for line in lines:
            if not line:
                if not previous_blank:
                    document.add_paragraph("")
                previous_blank = True
                continue
            add_paragraph(document, line)
            previous_blank = False
        if page_index != len(reader.pages) - 1:
            document.add_section(WD_SECTION.NEW_PAGE)

    document.save(str(DOCX_PATH))


if __name__ == "__main__":
    build_docx()
